import requests
import os
from docx import Document

# 1. Create two dummy DOCX files
doc1 = Document()
doc1.add_heading('DocIntel Test 1', 0)
doc1.add_paragraph('This is the first test document.')
doc1.save('test_sample1.docx')

doc2 = Document()
doc2.add_heading('DocIntel Test 2', 0)
doc2.add_paragraph('This is the second test document.')
doc2.save('test_sample2.docx')

# 2. Test the endpoint
URL = "http://localhost:8000/converters/word-to-pdf"

try:
    with open('test_sample1.docx', 'rb') as f1, open('test_sample2.docx', 'rb') as f2:
        files = [
            ('files', ('test_sample1.docx', f1, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')),
            ('files', ('test_sample2.docx', f2, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        print(f"Sending request to {URL}...")
        response = requests.post(URL, files=files)
        
        if response.status_code == 200:
            with open('test_result.pdf', 'wb') as out:
                out.write(response.content)
            print("✅ SUCCESS: Word to PDF converted! saved as test_result.pdf")
        else:
            print(f"❌ FAILED: Status {response.status_code}")
            print(f"Response: {response.text}")
finally:
    # Cleanup
    for p in ['test_sample1.docx', 'test_sample2.docx']:
        if os.path.exists(p):
            os.remove(p)
